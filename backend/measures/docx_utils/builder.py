"""
Утилиты сборки .docx в стиле исходного шаблона аудита.

Хелперы предполагают, что документ открыт из `template/base.docx` (стили и тема
подгружены автоматически), и делают типичные элементы раздела: заголовок,
подпись таблицы, абзац основного текста (с красной строкой), таблицу с рамками.

Форматирование повторяет оригинальный документ «Баян-Сулу»:
  • Шрифт: Times New Roman, 12pt.
  • Абзац: красная строка ≈ 1,25 см (450215 EMU), выравнивание по ширине,
    межстрочный интервал 1.0.
  • Заголовок раздела: жирный, 12pt, без красной строки, space_after = 12pt.
  • Подпись таблицы: жирный, без отступа, по левому краю.
  • Таблицы: Normal Table, рамки single 0.5pt, шрифт Times New Roman 12pt.
"""
from __future__ import annotations

from typing import Iterable, Sequence

from docx.document import Document as _DocumentType
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, Emu
from docx.enum.text import WD_COLOR_INDEX
import re


# Размер первой строки абзаца в исходнике — 450215 EMU ≈ 1,25 см.
BODY_FIRST_LINE_INDENT = Emu(450215)

# Шрифт и размер оригинала.
FONT_NAME = "Times New Roman"
FONT_SIZE = Pt(12)
TABLE_FONT_SIZE = Pt(10)


# --- абзацы -------------------------------------------------------------------

def _set_run_font(run, *, bold: bool = False, italic: bool = False) -> None:
    """Явно задаёт шрифт run-а, чтобы не зависеть от темы документа."""
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE
    run.bold = bold
    run.italic = italic

def add_runs_with_highlight(p, text: str, bold: bool = False, italic: bool = False) -> None:
    """Разбивает текст и подсвечивает номера (например 3.3.6.1) желтым маркером."""
    pattern = re.compile(r'(3\.3(?:\.\d+)+)')
    for part in pattern.split(text):
        if not part:
            continue
        run = p.add_run(part)
        _set_run_font(run, bold=bold, italic=italic)
        if pattern.match(part):
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def add_body_paragraph(
    doc: _DocumentType,
    text: str,
    *,
    bold: bool = False,
    align: WD_PARAGRAPH_ALIGNMENT = WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
    first_line_indent: bool = True,
) -> None:
    """Обычный абзац основного текста: красная строка, выравнивание по ширине."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.first_line_indent = BODY_FIRST_LINE_INDENT if first_line_indent else None
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)
    pf.line_spacing = 1.0
    p.alignment = align
    if text:
        add_runs_with_highlight(p, text, bold=bold)


def add_caption(doc: _DocumentType, text: str) -> None:
    """Подпись таблицы/рисунка: жирный, без отступа, по левому краю."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.first_line_indent = None
    pf.space_after = Pt(0)
    pf.space_before = Pt(6)
    pf.line_spacing = 1.0
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    add_runs_with_highlight(p, text, bold=True)


def add_heading_para(doc: _DocumentType, text: str) -> None:
    """Заголовок раздела «3.3.6 …»: жирный, без красной строки.

    В оригинале: space_after ≈ 12pt (152400 EMU), line_spacing 1.0.
    """
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.first_line_indent = None
    pf.space_after = Emu(152400)  # точно как в оригинале
    pf.space_before = Pt(0)
    pf.line_spacing = 1.0
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    add_runs_with_highlight(p, text, bold=True)


# --- таблицы ------------------------------------------------------------------

def _set_cell_background(cell, fill_color: str) -> None:
    """Устанавливает цвет фона ячейки таблицы (через XML)."""
    tcPr = cell._element.get_or_add_tcPr()
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tcPr.append(shd)
    
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_color)

def _set_xml_attr(elem, name: str, value: str) -> None:
    elem.set(qn(name), value)


def _make_border(kind: str, sz: int = 4) -> "OxmlElement":
    el = OxmlElement(f"w:{kind}")
    _set_xml_attr(el, "w:val", "single")
    _set_xml_attr(el, "w:sz", str(sz))
    _set_xml_attr(el, "w:space", "0")
    _set_xml_attr(el, "w:color", "auto")
    return el


def _apply_table_style(table) -> None:
    """Центровка таблицы, узкие поля ячеек и сплошные рамки 0,5pt — как в исходнике."""
    tbl = table._element
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)

    # Центровка
    jc = tblPr.find(qn("w:jc"))
    if jc is None:
        jc = OxmlElement("w:jc")
        tblPr.append(jc)
    _set_xml_attr(jc, "w:val", "center")

    # Узкие поля ячеек (70 twips ≈ 1.23 мм со всех сторон)
    cellMar = tblPr.find(qn("w:tblCellMar"))
    if cellMar is None:
        cellMar = OxmlElement("w:tblCellMar")
        tblPr.append(cellMar)
    for side in ("top", "left", "bottom", "right"):
        node = cellMar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            cellMar.append(node)
        _set_xml_attr(node, "w:w", "70")
        _set_xml_attr(node, "w:type", "dxa")

    # Сплошные рамки.
    borders = tblPr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblPr.append(borders)
    for kind in ("top", "left", "bottom", "right", "insideH", "insideV"):
        old = borders.find(qn(f"w:{kind}"))
        if old is not None:
            borders.remove(old)
        borders.append(_make_border(kind))


def _set_cell_font(cell) -> None:
    """Принудительно задаём Times New Roman 10pt для всех run-ов в ячейке."""
    for p in cell.paragraphs:
        pf = p.paragraph_format
        pf.space_before = Pt(1)
        pf.space_after = Pt(1)
        pf.line_spacing = 1.0
        for run in p.runs:
            run.font.name = FONT_NAME
            run.font.size = TABLE_FONT_SIZE


def make_audit_table(
    doc: _DocumentType,
    header: Sequence[str],
    rows: Iterable[Sequence[str]],
    *,
    numeric_columns: Sequence[int] | None = None,
    col_widths: Sequence[Any] | None = None,  # <--- ДОБАВЛЕН НОВЫЙ АРГУМЕНТ
) -> None:
    """Таблица в стиле шаблона аудита.

    - Стиль `Normal Table` (унаследован из base.docx).
    - Центровка таблицы, узкие поля ячеек, видимые границы.
    - Шапка: по центру, vertical-center. Шрифт Times New Roman 12pt.
    - Колонки в `numeric_columns` выровнены по правому краю.
    """
    numeric_columns = set(numeric_columns or [])
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Normal Table"
    table.autofit = False
    hdr_cells = table.rows[0].cells
    for idx, value in enumerate(header):
        cell = hdr_cells[idx]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(value)
        run.font.name = FONT_NAME
        run.font.size = TABLE_FONT_SIZE
        run.bold = False  # В оригинале шапки не жирные (кроме Table 4)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _set_cell_background(cell, "D9E2F3")

    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cell = cells[idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = (
                WD_PARAGRAPH_ALIGNMENT.RIGHT
                if idx in numeric_columns
                else WD_PARAGRAPH_ALIGNMENT.LEFT
            )
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.line_spacing = 1.0
            run = p.add_run(str(value))
            run.font.name = FONT_NAME
            run.font.size = TABLE_FONT_SIZE
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
# Применяем заданную ширину к каждой ячейке
    if col_widths:
        for row in table.rows:
            for idx, width in enumerate(col_widths):
                if width is not None and idx < len(row.cells):
                    row.cells[idx].width = width
    # ------------------------------------
    _apply_table_style(table)
