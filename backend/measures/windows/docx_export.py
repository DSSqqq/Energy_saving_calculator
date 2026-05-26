"""
Сборка раздела отчёта по мероприятию «Окна» в формате .docx.

Структура повторяет исходный шаблон аудита (Баян-Сулу, раздел 3.3.6):
  1. Таблица 3.3.6.1 — площадь проблемных участков по зданиям.
  2. Поясняющий абзац.
  3. Таблица 3.3.6.2 — исходные данные (общие + per-building период/температура).
  4. На каждое здание: блок расчёта Q_т (теплопередача).
  5. На каждое здание: блок расчёта Q_inf (инфильтрация).
  6. Сумма по зданиям, перевод в м³ газа, экономия в тг, инвестиции, простой PBP.
  7. Таблица 3.3.6.3 — материалы и стоимость.
  8. Таблица 3.3.6.4 — экономия по зданиям.
  9. Таблица 3.3.1.1 — финансовые показатели проекта.
 10. Диаграмма 3.3.1.1 — кумулятивная дисконтированная NPV (matplotlib → PNG).

Документ открывается из заранее подготовленной базы `template/base.docx`,
которая содержит стили, тему и шрифты исходного шаблона клиента — все добавляемые
программно элементы наследуют те же фонты и оформление. Формулы вставляются как
настоящие Word-уравнения (OMML, Cambria Math italic) через `equations.eq_*`.
"""
from __future__ import annotations

import io
from pathlib import Path
from typing import Any, Dict

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt, Cm

from ..docx_utils.builder import (
    add_body_paragraph,
    add_caption,
    add_heading_para,
    make_audit_table,
    add_runs_with_highlight,
)
from ..docx_utils.chart import render_npv_chart
from ..docx_utils.number_fmt import (
    fmt_money,
    fmt_num,
    fmt_percent,
    fmt_years,
)
from . import equations as eq

from docx.enum.section import WD_SECTION, WD_ORIENT
from docx.shared import Cm, Pt, Inches

BASE_DOCX_PATH = Path(__file__).parent / "template" / "base.docx"


_INTRO_AFTER_TABLE_1 = (
    "Таким образом, выявленные дефекты представляют собой системную проблему, "
    "требующую комплексного подхода к устранению. Предлагается устранить щели "
    "и неплотности, выполнить ремонтные работы по устранению дефектов и утеплению."
)
_INTRO_BEFORE_INPUTS = (
    "Ниже приведён расчёт, выполненный с учётом того, что из общего объёма "
    "зафиксированных окон в расчёт принята наиболее проблемная доля. Предлагаемые "
    "мероприятия повышают сопротивление теплопередаче и снижают воздухопроницаемость "
    "оконных блоков."
)

def build_docx(payload: Dict[str, Any], results: Dict[str, Any]) -> bytes:
    """Сборка документа программно из чистого base.docx (стили + размеры страницы)."""
    if BASE_DOCX_PATH.exists():
        doc = Document(str(BASE_DOCX_PATH))
    else:
        doc = Document()
        _apply_fallback_styles(doc)

    add_heading_para(doc, "3.3.6 Улучшение теплозащитных свойств оконных блоков")

    _section_table_1(doc, results)
    add_body_paragraph(doc, _INTRO_AFTER_TABLE_1)
    add_body_paragraph(doc, _INTRO_BEFORE_INPUTS)

    if results["buildings"]:
        first_building = results["buildings"][0]
        _section_per_building_transmission(doc, first_building)
        _section_per_building_infiltration(doc, first_building, payload["shared"])
        
        if len(results["buildings"]) > 1:
            add_body_paragraph(doc, "Алгоритм расчета для остальных сооружений идентичен вышеописанному (см. таблицу 3.3.6.4).")
            doc.add_paragraph()

    _section_totals(doc, payload, results)
    _section_table_3(doc, results)
    _section_table_4(doc, results)
    _section_table_finance(doc, results)
    _section_chart(doc, results)

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


# --- helpers ------------------------------------------------------------------

def _apply_fallback_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)

    pf = normal.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.0

def _add_equation(doc, omml_element) -> None:
    """Вставляет готовый OMML-абзац в документ и центрирует его."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    
    p = doc.add_paragraph()
    
    # Обнуляем отступы, чтобы они не мешали центровке
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.left_indent = Pt(0)
    p.paragraph_format.right_indent = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(2)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Центрируем саму формулу внутри OMML
    if omml_element.tag.endswith('oMathPara'):
        m_pr = omml_element.find(qn('m:oMathParaPr'))
        if m_pr is None:
            m_pr = OxmlElement('m:oMathParaPr')
            omml_element.insert(0, m_pr)
        m_jc = m_pr.find(qn('m:jc'))
        if m_jc is None:
            m_jc = OxmlElement('m:jc')
            m_pr.append(m_jc)
        m_jc.set(qn('m:val'), 'centerGroup')

    p._p.append(omml_element)

def _bullet_paragraph(doc, text: str) -> None:
    """Маркеро-подобная строка в исходных данных (с родным буллетом Word)."""
    p = doc.add_paragraph(style="List Bullet")
    
    p.paragraph_format.left_indent = Cm(1.89)
    p.paragraph_format.first_line_indent = Cm(-0.63)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def _bullet_rich(doc, *segments) -> None:
    """Маркированная строка с подстрочными индексами.

    Каждый segment — либо str (обычный текст), либо tuple (base, sub),
    где base — текст перед индексом, sub — подстрочный текст.
    Пример: _bullet_rich(doc, "Средняя наружная температура: ", ("t", "н"), " = 20 °С")
    """
    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement

    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(1.89)
    p.paragraph_format.first_line_indent = Cm(-0.63)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    def _add_run(text, subscript=False):
        run = p.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        if subscript:
            rpr = run._element.get_or_add_rPr()
            va = OxmlElement("w:vertAlign")
            va.set(_qn("w:val"), "subscript")
            rpr.append(va)

    for seg in segments:
        if isinstance(seg, str):
            _add_run(seg)
        elif isinstance(seg, tuple) and len(seg) == 2:
            base, sub = seg
            _add_run(base)
            _add_run(sub, subscript=True)
        else:
            _add_run(str(seg))


# --- секции -------------------------------------------------------------------

def _section_table_1(doc, results: Dict[str, Any]) -> None:
    from docx.enum.section import WD_SECTION, WD_ORIENT

    # --- ПЕРЕВОДИМ СТРАНИЦУ В АЛЬБОМНУЮ ОРИЕНТАЦИЮ ---
    landscape_section = doc.add_section(WD_SECTION.NEW_PAGE)
    landscape_section.orientation = WD_ORIENT.LANDSCAPE
    new_width, new_height = landscape_section.page_height, landscape_section.page_width
    landscape_section.page_width = new_width
    landscape_section.page_height = new_height

    # --- РУЧНОЕ СОЗДАНИЕ ПОДПИСИ ТАБЛИЦЫ ---
    p_cap = doc.add_paragraph()
    p_cap.paragraph_format.first_line_indent = None
    p_cap.paragraph_format.space_before = Pt(6)
    p_cap.paragraph_format.space_after = Pt(0)
    p_cap.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    
    add_runs_with_highlight(p_cap, "Таблица 3.3.6.1 - Исходные данные", italic=True)
    # ---------------------------------------
    rows = []
    for b in results["buildings"]:
        rows.append((
            b["name"],
            fmt_num(b["area_m2"], 2),
            fmt_num(b["r_before"], 2),
            fmt_num(b["r_after"], 2),
            fmt_num(b["period_days"], 0),
            fmt_num(b["t_inside"], 1),
            fmt_num(b["t_outside_avg"], 1),
        ))
    
    rows.append((
        "ИТОГО",
        fmt_num(results["totals"]["area_m2"], 2),
        "-", "-", "-", "-", "-"
    ))
    
    headers = [
        "Наименование здания",
        "Длина проблемных участков, м.п.",
        "Фактическое сопротивление теплопередаче проблемных участков до мероприятия, м.п.·°С/Вт",
        "Сопротивление теплопередаче проблемных участков после мероприятия, м.п.·°С/Вт",
        "Продолжительность отопительного периода, сут.",
        "Температура внутри помещения, °С",
        "Средняя температура наружного воздуха за отопительный период, °С"
    ]
    
    make_audit_table(
        doc,
        headers,
        rows,
        numeric_columns=[1, 2, 3, 4, 5, 6],
        col_widths=[Cm(3), Cm(3), Cm(4), Cm(4), Cm(3), Cm(3), Cm(4)]
    )
    doc.add_paragraph() # <--- Пустая строка
    
    # --- ВОЗВРАЩАЕМ КНИЖНУЮ ОРИЕНТАЦИЮ ДЛЯ СЛЕДУЮЩИХ РАЗДЕЛОВ ---
    portrait_section = doc.add_section(WD_SECTION.NEW_PAGE)
    portrait_section.orientation = WD_ORIENT.PORTRAIT
    p_width, p_height = portrait_section.page_height, portrait_section.page_width
    portrait_section.page_width = p_width
    portrait_section.page_height = p_height


def _section_per_building_transmission(doc, b: Dict[str, Any]) -> None:
    add_body_paragraph(
        doc,
        f"Ниже представлен расчёт теплопотерь через дефекты оконных блоков для здания {b['name']}.",
    )
    add_body_paragraph(doc, "Исходные данные:", first_line_indent=False)
    _bullet_paragraph(doc, f"Длина проблемных участков: L = {fmt_num(b['area_m2'], 2)} м.п.")
    _bullet_rich(doc, "Средняя наружная температура: ", ("t", "н"), f" = {fmt_num(b['t_outside_avg'], 1)} °С")
    _bullet_rich(doc, "Внутренняя температура: ", ("t", "в"), f" = {fmt_num(b['t_inside'], 1)} °С")
    _bullet_paragraph(doc, f"Продолжительность отопительного периода: z = {fmt_num(b['period_days'], 0)} сут.")
    _bullet_rich(doc, "Сопротивление теплопередаче до мероприятия: ", ("R", "до"), f" = {fmt_num(b['r_before'], 2)} м.п.·°С/Вт")
    _bullet_rich(doc, "Сопротивление теплопередаче после мероприятия: ", ("R", "после"), f" = {fmt_num(b['r_after'], 2)} м.п.·°С/Вт")
    add_body_paragraph(doc, "Теплопотери до мероприятия:")
    doc.add_paragraph() # <--- Пустая строка
    _add_equation(
        doc,
        eq.eq_q_kwh(
            subscript="1",
            r_subscript="до",
            area=fmt_num(b["area_m2"], 2),
            r_val=fmt_num(b["r_before"], 2),
            t_in=fmt_num(b["t_inside"], 1),
            t_out_avg=fmt_num(b["t_outside_avg"], 1),
            period_days=fmt_num(b["period_days"], 0),
            result_kwh=fmt_num(b["q1_kwh"], 1),
        ),
    )
    doc.add_paragraph() # <--- Пустая строка
    add_body_paragraph(doc, "Теплопотери после мероприятия:")
    doc.add_paragraph() # <--- Пустая строка
    _add_equation(
        doc,
        eq.eq_q_kwh(
            subscript="2",
            r_subscript="после",
            area=fmt_num(b["area_m2"], 2),
            r_val=fmt_num(b["r_after"], 2),
            t_in=fmt_num(b["t_inside"], 1),
            t_out_avg=fmt_num(b["t_outside_avg"], 1),
            period_days=fmt_num(b["period_days"], 0),
            result_kwh=fmt_num(b["q2_kwh"], 1),
        ),
    )
    doc.add_paragraph() # <--- Пустая строка
    add_body_paragraph(doc, "Экономия тепловой энергии, Гкал:")
    doc.add_paragraph() # <--- Пустая строка
    _add_equation(
        doc,
        eq.eq_delta_q(
            q1=fmt_num(b["q1_kwh"], 1),
            q2=fmt_num(b["q2_kwh"], 1),
            result_gcal=fmt_num(b["q_transmission_gcal"], 2),
        ),
    )
    doc.add_paragraph() # <--- Пустая строка


def _section_per_building_infiltration(doc, b: Dict[str, Any], shared: Dict[str, Any]) -> None:
    add_body_paragraph(
        doc,
        f"Расчёт удельного расхода теплоты на нагревание инфильтрующегося воздуха через "
        f"оконные блоки для здания {b['name']}, ккал/(м.п.·ч):",
    )
    add_body_paragraph(doc, "Исходные данные:", first_line_indent=False)
    _bullet_rich(
        doc,
        ("g", "inf1"),
        f" = {fmt_num(b['g_inf_before'], 2)} кг/(м.п.·ч), коэффициент воздухопроницаемости существующих окон;",
    )
    _bullet_rich(
        doc,
        ("g", "inf2"),
        f" = {fmt_num(b['g_inf_after'], 2)} кг/(м.п.·ч), коэффициент воздухопроницаемости устанавливаемых окон;",
    )
    _bullet_paragraph(
        doc,
        f"c = {fmt_num(shared['c_air'], 2)} ккал/(кг·°С), удельная теплоёмкость воздуха;",
    )
    _bullet_paragraph(
        doc,
        f"k = {fmt_num(shared['k_factor'], 2)}, коэффициент влияния встречного теплового потока.",
    )
    add_body_paragraph(doc, "До мероприятия, ккал/(м.п.·ч):", first_line_indent=False)
    doc.add_paragraph() # <--- Пустая строка
    _add_equation(
        doc,
        eq.eq_q_inf_specific(
            label_sub="до",
            g_inf_label="inf1",
            c_air=fmt_num(shared["c_air"], 2),
            g_inf=fmt_num(b["g_inf_before"], 2),
            k_factor=fmt_num(shared["k_factor"], 2),
            delta_t=fmt_num(b["delta_t"], 2),
            result_kcal=fmt_num(b["q_inf_specific_before"], 2),
        ),
    )
    doc.add_paragraph() # <--- Пустая строка
    add_body_paragraph(doc, "После мероприятия, ккал/(м.п.·ч):", first_line_indent=False)
    doc.add_paragraph() # <--- Пустая строка
    _add_equation(
        doc,
        eq.eq_q_inf_specific(
            label_sub="после",
            g_inf_label="inf2",
            c_air=fmt_num(shared["c_air"], 2),
            g_inf=fmt_num(b["g_inf_after"], 2),
            k_factor=fmt_num(shared["k_factor"], 2),
            delta_t=fmt_num(b["delta_t"], 2),
            result_kcal=fmt_num(b["q_inf_specific_after"], 2),
        ),
    )
    doc.add_paragraph() # <--- Пустая строка
    add_body_paragraph(doc, "Объём сэкономленной тепловой энергии:", first_line_indent=False)
    doc.add_paragraph() # <--- Пустая строка
    _add_equation(
        doc,
        eq.eq_q_infiltration(
            q_before=fmt_num(b["q_inf_specific_before"], 2),
            q_after=fmt_num(b["q_inf_specific_after"], 2),
            area=fmt_num(b["area_m2"], 2),
            period_days=fmt_num(b["period_days"], 0),
            result_gcal=fmt_num(b["q_infiltration_gcal"], 2),
        ),
    )
    doc.add_paragraph() # <--- Пустая строка


def _section_totals(doc, payload: Dict[str, Any], results: Dict[str, Any]) -> None:
    add_body_paragraph(
        doc,
        "Общая экономия тепловой энергии за счёт повышения уровня теплозащиты окон:",
    )
    for b in results["buildings"]:
        _add_equation(
            doc,
            eq.eq_sum_two(
                label=f"Здание {b['name']}",
                q_t=fmt_num(b["q_transmission_gcal"], 2),
                q_inf=fmt_num(b["q_infiltration_gcal"], 2),
                total=fmt_num(b["q_total_gcal"], 2),
            ),
        )
    totals = results["totals"]
    add_body_paragraph(doc, f"ИТОГО: ΣQ = {fmt_num(totals['q_total_gcal'], 2)} Гкал", first_line_indent=False)
    
    FUEL_META = {
        "gas": {"label": "природного газа", "label_nom": "Природный газ", "label_short": "газа", "fuel_unit": "тыс. м³", "tariff_unit": "тг/м³", "cv_unit": "Гкал/тыс. м³", "multiplier": 1000},
        "electricity": {"label": "электрической энергии", "label_nom": "Электрическая энергия", "label_short": "электр.", "fuel_unit": "тыс. кВт·ч", "tariff_unit": "тг/кВт·ч", "cv_unit": "Гкал/тыс. кВт·ч", "multiplier": 1000},
        "coal": {"label": "каменного угля", "label_nom": "Каменный уголь", "label_short": "угля", "fuel_unit": "тонн", "tariff_unit": "тг/тонна", "cv_unit": "Гкал/тонна", "multiplier": 1},
        "diesel": {"label": "дизельного топлива", "label_nom": "Дизельное топливо", "label_short": "диз.", "fuel_unit": "тонн", "tariff_unit": "тг/тонна", "cv_unit": "Гкал/тонна", "multiplier": 1},
        "gcal": {"label": "тепловой энергии", "label_nom": "Центральное теплоснабжение", "label_short": "тепла", "fuel_unit": "Гкал", "tariff_unit": "тг/Гкал", "cv_unit": "Гкал/Гкал", "multiplier": 1},
    }
    
    first_b = results["buildings"][0] if results["buildings"] else {}
    first_fuel = first_b.get("fuel_type", "gas")
    all_same_fuel = all(row.get("fuel_type") == first_fuel for row in results["buildings"])

    if all_same_fuel:
        fuel_type = first_fuel
        fuel_calorific = first_b.get("fuel_calorific", 1.0)
        fuel_tariff = first_b.get("fuel_tariff", 0.0)
        meta = FUEL_META.get(fuel_type, FUEL_META["gas"])

        if fuel_type != "gcal":
            add_body_paragraph(
                doc,
                f"Перевод потребленных Гкал в объем {meta['label']} ({meta['fuel_unit']}); удельная теплота сгорания принята за "
                f"{fmt_num(fuel_calorific, 4)} {meta['cv_unit']}:",
            )
            doc.add_paragraph() # <--- Пустая строка
            _add_equation(
                doc,
                eq.eq_fuel_savings(
                    label=meta["label_short"],
                    sum_q=fmt_num(totals["q_total_gcal"], 2),
                    calorific=fmt_num(fuel_calorific, 4),
                    result=fmt_num(totals["fuel_savings"], 3),
                    unit=meta["fuel_unit"],
                ),
            )
            doc.add_paragraph() # <--- Пустая строка

        add_body_paragraph(doc, "Экономия тепловой энергии в денежном выражении, тг:")
        doc.add_paragraph() # <--- Пустая строка
        
        if fuel_type != "gcal":
            _add_equation(
                doc,
                eq.eq_money_savings(
                    fuel_val=fmt_num(totals["fuel_savings"], 3),
                    tariff=fmt_num(fuel_tariff, 2),
                    multiplier=str(meta["multiplier"]),
                    result=fmt_money(totals["money_savings_tg"]),
                ),
            )
            doc.add_paragraph() # <--- Пустая строка
            add_body_paragraph(
                doc,
                f"где Cтэ = {fmt_num(fuel_tariff, 2)} {meta['tariff_unit']} — средняя цена {meta['label']}.",
            )
        else:
            _add_equation(
                doc,
                eq.eq_money_savings(
                    fuel_val=fmt_num(totals["q_total_gcal"], 2),
                    tariff=fmt_num(fuel_tariff, 2),
                    multiplier="1",
                    result=fmt_money(totals["money_savings_tg"]),
                ),
            )
            doc.add_paragraph() # <--- Пустая строка
            add_body_paragraph(
                doc,
                f"где Cтэ = {fmt_num(fuel_tariff, 2)} тг/Гкал — средняя цена тепловой энергии.",
            )
    else:
        add_body_paragraph(
            doc,
            "В связи с тем, что здания отапливаются от различных источников энергии (видов топлива), перевод в объемы замещаемого энергоносителя и расчет денежной экономии производится индивидуально для каждого здания:",
        )
        doc.add_paragraph()
        
        # --- РУЧНОЕ СОЗДАНИЕ ПОДПИСИ ТАБЛИЦЫ ---
        p_cap = doc.add_paragraph()
        p_cap.paragraph_format.first_line_indent = None
        p_cap.paragraph_format.space_before = Pt(6)
        p_cap.paragraph_format.space_after = Pt(0)
        p_cap.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        
        add_runs_with_highlight(
            p_cap,
            "Таблица 3.3.6.2 - Параметры индивидуального теплоснабжения и экономии",
            italic=True
        )
        # ---------------------------------------
        
        tbl_headers = [
            "Наименование здания",
            "Источник отопления",
            "Теплотворность",
            "Тариф за ед.",
            "Экономия топлива",
            "Экономия в год, тг"
        ]
        tbl_rows = []
        for b in results["buildings"]:
            b_meta = FUEL_META.get(b["fuel_type"], FUEL_META["gas"])
            cv_str = f"{fmt_num(b['fuel_calorific'], 4)} {b_meta['cv_unit']}" if b["fuel_type"] != "gcal" else "-"
            tbl_rows.append((
                b["name"],
                b_meta["label_nom"],
                cv_str,
                f"{fmt_num(b['fuel_tariff'], 2)} {b_meta['tariff_unit']}",
                f"{fmt_num(b['fuel_savings'], 2)} {b_meta['fuel_unit']}",
                fmt_money(b["money_savings_tg"])
            ))
        
        make_audit_table(
            doc,
            tbl_headers,
            tbl_rows,
            numeric_columns=[2, 3, 4, 5],
            col_widths=[Cm(3.5), Cm(3.5), Cm(3.5), Cm(3), Cm(3), Cm(3.5)]
        )
        doc.add_paragraph()
    invest = results["investment"]
    add_body_paragraph(doc, "Стоимость затрат на внедрение мероприятия, тг:")
    add_body_paragraph(
        doc,
        f"I = {fmt_money(invest['total_tg'])} тг (см. таблицу 3.3.6.3).",
        first_line_indent=False,
    )
    add_body_paragraph(doc, "Простой срок окупаемости мероприятия, год:")
    doc.add_paragraph() # <--- Пустая строка
    _add_equation(
        doc,
        eq.eq_pbp(
            investment=fmt_money(invest["total_tg"]),
            money=fmt_money(totals["money_savings_tg"]),
            result=fmt_years(results["finance"]["pbp_years"]),
        ),
    )
    doc.add_page_break() # <--- Разрыв страницы


def _section_table_3(doc, results: Dict[str, Any]) -> None:
    doc.add_paragraph() # <--- Пустая строка
    # --- РУЧНОЕ СОЗДАНИЕ ПОДПИСИ ТАБЛИЦЫ ---
    p_cap = doc.add_paragraph()
    p_cap.paragraph_format.first_line_indent = None
    p_cap.paragraph_format.space_before = Pt(6)
    p_cap.paragraph_format.space_after = Pt(0)
    p_cap.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    
    add_runs_with_highlight(p_cap, "Таблица 3.3.6.3 - Материалы и ориентировочная стоимость утепления", italic=True)
    # ---------------------------------------
    rows = []
    for item in results["investment"]["items"]:
        rows.append(
            (
                item["name"],
                fmt_num(item["quantity"], 2),
                fmt_num(item["price_per_unit"], 2),
                fmt_money(item["cost_tg"]),
            )
        )
    rows.append(("Итого", "—", "—", fmt_money(results['investment']['total_tg'])))
    make_audit_table(
        doc,
        ["Материал", "Количество, м.п.", "Цена за ед., тг", "Стоимость, тг"],
        rows,
        numeric_columns=[1, 2, 3], col_widths=[Cm(6), Cm(3), Cm(4), Cm(4)]
    )
    doc.add_paragraph() # <--- Пустая строка

def _section_table_4(doc, results: Dict[str, Any]) -> None:
# --- РУЧНОЕ СОЗДАНИЕ ПОДПИСИ ТАБЛИЦЫ ---
    p_cap = doc.add_paragraph()
    p_cap.paragraph_format.first_line_indent = None
    p_cap.paragraph_format.space_before = Pt(6)
    p_cap.paragraph_format.space_after = Pt(0)
    p_cap.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    
    add_runs_with_highlight(p_cap, "Таблица 3.3.6.4 – Экономия тепловой энергии в Гкал и денежном эквиваленте", italic=True)
    # ---------------------------------------
    rows = []
    for b in results["buildings"]:
        rows.append(
            (
                b["name"],
                fmt_num(b["q_total_gcal"], 2),
                fmt_money(b["money_savings_tg"]),
            )
        )
    rows.append(
        (
            "ИТОГО",
            fmt_num(results["totals"]["q_total_gcal"], 2),
            fmt_money(results["totals"]["money_savings_tg"]),
        )
    )
    make_audit_table(
        doc,
        [
            "Наименование здания",
            "Экономия энергии, Гкал",
            "Экономия в денежном выражении, тг",
        ],
        rows,
        numeric_columns=[1, 2], col_widths=[Cm(6), Cm(5), Cm(6)]
    )
    
    doc.add_paragraph() # <--- Пустая строка

def _section_table_finance(doc, results: Dict[str, Any]) -> None:
# --- РУЧНОЕ СОЗДАНИЕ ПОДПИСИ ТАБЛИЦЫ ---
    p_cap = doc.add_paragraph()
    p_cap.paragraph_format.first_line_indent = None
    p_cap.paragraph_format.space_before = Pt(6)
    p_cap.paragraph_format.space_after = Pt(0)
    p_cap.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    
    add_runs_with_highlight(p_cap, "Таблица 3.3.1.1 - Показатели эффективности проекта", italic=True)
    # ---------------------------------------
    f = results["finance"]
    rows = [
        ("Ставка дисконтирования", "(%)", fmt_num(f["discount_rate"] * 100.0, 2)),
        ("Первоначальные инвестиции", "тг", fmt_money(results["investment"]["total_tg"])),
        ("Чистая приведённая стоимость (NPV)", "тг", fmt_money(f["npv_tg"])),
        ("Дисконтированный индекс доходности (DPI)", "о.е.", fmt_num(f["dpi"], 2)),
        (
            "Внутренняя норма доходности (IRR)",
            "(%)",
            fmt_percent(f["irr"]) if f["irr"] is not None else "—",
        ),
        ("Простой срок окупаемости (PBP)", "лет", fmt_years(f["pbp_years"])),
        ("Дисконтированный срок окупаемости (DPBP)", "год", fmt_years(f["dpbp_years"])),
    ]
    make_audit_table(doc, ["Наименование", "Ед. изм-я", "Значение"], rows, 
    numeric_columns=[2], col_widths=[Cm(11), Cm(3), Cm(3)])


def _section_chart(doc, results: Dict[str, Any]) -> None:
    # --- СОЗДАЕМ НОВЫЙ РАЗДЕЛ С НОВОЙ СТРАНИЦЫ ---
    new_section = doc.add_section(WD_SECTION.NEW_PAGE)
    
    # Меняем ориентацию на альбомную
    new_section.orientation = WD_ORIENT.LANDSCAPE
    
    # Меняем местами ширину и высоту страницы (обязательный шаг для python-docx)
    new_width, new_height = new_section.page_height, new_section.page_width
    new_section.page_width = new_width
    new_section.page_height = new_height
    # ---------------------------------------------

# Генерируем картинку
    png = render_npv_chart(results["finance"]["npv_by_year"])
    # --- 1. СНАЧАЛА СОЗДАЕМ АБЗАЦ ДЛЯ КАРТИНКИ ---
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.paragraph_format.first_line_indent = None
    # Убираем отступ после картинки, чтобы подпись была ближе к ней (по желанию)
    p.paragraph_format.space_after = Pt(0) 
    run = p.add_run()
    # Вставляем картинку.
    run.add_picture(io.BytesIO(png), width=Inches(9.0))
    # --- 2. ЗАТЕМ СОЗДАЕМ ПОДПИСЬ ПОД КАРТИНКОЙ ---
    p_cap = doc.add_paragraph()
    p_cap.paragraph_format.first_line_indent = None
    p_cap.paragraph_format.space_before = Pt(6) # Небольшой отступ между графиком и подписью
    p_cap.paragraph_format.space_after = Pt(12)
    p_cap.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    add_runs_with_highlight(p_cap, "Диаграмма 3.3.1.1 – Чистая приведённая стоимость проекта мероприятия по "
        "улучшению теплозащитных свойств оконных блоков", italic=True)
    # -----------------------------------------------
