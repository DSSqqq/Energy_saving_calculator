"""
Сборка раздела отчёта по мероприятию «Стены» в формате .docx.
Логика разметки и стилей полностью совпадает с мероприятием «Окна».
"""
from __future__ import annotations

import io
from pathlib import Path
from typing import Any, Dict

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt, Cm
from docx.enum.section import WD_SECTION, WD_ORIENT

from ..docx_utils.builder import (
    add_body_paragraph,
    add_caption,
    add_heading_para,
    make_audit_table,
    add_runs_with_highlight,
)
from ..docx_utils.chart import render_npv_chart
from ..docx_utils.number_fmt import (
    fmt_money,
    fmt_num,
    fmt_percent,
    fmt_years,
)
from ..windows import equations as eq

BASE_DOCX_PATH = Path(__file__).parent.parent / "windows" / "template" / "base.docx"

FUEL_META = {
    "gas": {"label": "природного газа", "label_nom": "Природный газ", "label_short": "газа", "fuel_unit": "тыс. м³", "tariff_unit": "тг/м³", "cv_unit": "Гкал/тыс. м³", "multiplier": "1000"},
    "electricity": {"label": "электрической энергии", "label_nom": "Электрическая энергия", "label_short": "электр.", "fuel_unit": "тыс. кВт·ч", "tariff_unit": "тг/кВт·ч", "cv_unit": "Гкал/тыс. кВт·ч", "multiplier": "1000"},
    "coal": {"label": "каменного угля", "label_nom": "Каменный уголь", "label_short": "угля", "fuel_unit": "тонн", "tariff_unit": "тг/тонна", "cv_unit": "Гкал/тонна", "multiplier": "1"},
    "diesel": {"label": "дизельного топлива", "label_nom": "Дизельное топливо", "label_short": "диз.", "fuel_unit": "тонн", "tariff_unit": "тг/тонна", "cv_unit": "Гкал/тонна", "multiplier": "1"},
    "gcal": {"label": "тепловой энергии", "label_nom": "Центральное теплоснабжение", "label_short": "тепла", "fuel_unit": "Гкал", "tariff_unit": "тг/Гкал", "cv_unit": "Гкал/Гкал", "multiplier": "1"},
}

_INTRO_AFTER_TABLE_1 = (
    "Анализ наружных стен выявил высокий потенциал сбережения тепловой энергии за счет "
    "повышения их теплозащитных свойств (утепления). Предлагается выполнить комплексные "
    "работы по монтажу теплоизоляционного слоя на внешние фасады сооружений."
)
_INTRO_BEFORE_INPUTS = (
    "Ниже представлен детальный расчёт сбережения тепловой энергии. Внедрение теплоизоляции "
    "позволяет существенно снизить коэффициент теплопередачи наружных стен и повысить их термическое "
    "сопротивление, тем самым сокращая трансмиссионные потери теплоты."
)


def build_docx(payload: Dict[str, Any], results: Dict[str, Any]) -> bytes:
    """Сборка документа программно из base.docx (стили + размеры страницы)."""
    if BASE_DOCX_PATH.exists():
        doc = Document(str(BASE_DOCX_PATH))
    else:
        doc = Document()
        _apply_fallback_styles(doc)

    add_heading_para(doc, "3.3.7 Улучшение теплозащитных свойств наружных стен")

    _section_table_1(doc, results)
    add_body_paragraph(doc, _INTRO_AFTER_TABLE_1)
    add_body_paragraph(doc, _INTRO_BEFORE_INPUTS)

    if results["buildings"]:
        first_building = results["buildings"][0]
        _section_per_building_transmission(doc, first_building)
        
        if len(results["buildings"]) > 1:
            add_body_paragraph(doc, "Алгоритм расчета для остальных сооружений идентичен вышеописанному (см. таблицу 3.3.7.3).")
            doc.add_paragraph()

    _section_totals(doc, payload, results)
    _section_table_3(doc, results)
    _section_table_4(doc, results)
    _section_table_finance(doc, results)
    _section_chart(doc, results)

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


# --- Helpers ------------------------------------------------------------------

def _apply_fallback_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)

    pf = normal.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.0


def _add_equation(doc, omml_element) -> None:
    """Вставляет готовый OMML-абзац в документ и центрирует его."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.left_indent = Pt(0)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    m_pr = omml_element.find(qn('m:oMathParaPr'))
    if m_pr is None:
        m_pr = OxmlElement('m:oMathParaPr')
        omml_element.insert(0, m_pr)
    m_jc = m_pr.find(qn('m:jc'))
    if m_jc is None:
        m_jc = OxmlElement('m:jc')
        m_pr.append(m_jc)
    m_jc.set(qn('m:val'), 'centerGroup')

    p._p.append(omml_element)


def _bullet_paragraph(doc, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(1.89)
    p.paragraph_format.first_line_indent = Cm(-0.63)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def _bullet_rich(doc, *segments) -> None:
    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement

    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(1.89)
    p.paragraph_format.first_line_indent = Cm(-0.63)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    def _add_run(text, subscript=False):
        run = p.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        if subscript:
            rpr = run._element.get_or_add_rPr()
            va = OxmlElement("w:vertAlign")
            va.set(_qn("w:val"), "subscript")
            rpr.append(va)

    for seg in segments:
        if isinstance(seg, str):
            _add_run(seg)
        elif isinstance(seg, tuple) and len(seg) == 2:
            base, sub = seg
            _add_run(base)
            _add_run(sub, subscript=True)
        else:
            _add_run(str(seg))


# --- Sections -----------------------------------------------------------------

def _section_table_1(doc, results: Dict[str, Any]) -> None:
    landscape_section = doc.add_section(WD_SECTION.NEW_PAGE)
    landscape_section.orientation = WD_ORIENT.LANDSCAPE
    new_width, new_height = landscape_section.page_height, landscape_section.page_width
    landscape_section.page_width = new_width
    landscape_section.page_height = new_height

    p_cap = doc.add_paragraph()
    p_cap.paragraph_format.first_line_indent = None
    p_cap.paragraph_format.space_before = Pt(6)
    p_cap.paragraph_format.space_after = Pt(0)
    p_cap.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    
    add_runs_with_highlight(p_cap, "Таблица 3.3.7.1 - Исходные данные по зданиям", italic=True)

    rows = []
    for b in results["buildings"]:
        rows.append((
            b["name"],
            fmt_num(b["area_m2"], 2),
            fmt_num(b["r_before"], 2),
            fmt_num(b["r_after"], 2),
            fmt_num(b["period_days"], 0),
            fmt_num(b["t_inside"], 1),
            fmt_num(b["t_outside_avg"], 1),
        ))
    
    rows.append((
        "ИТОГО",
        fmt_num(results["totals"]["area_m2"], 2),
        "-", "-", "-", "-", "-"
    ))
    
    headers = [
        "Наименование здания",
        "Площадь наружных стен, м²",
        "Термическое сопротивление до, м²·°С/Вт",
        "Термическое сопротивление после, м²·°С/Вт",
        "Продолжительность отопительного периода, сут.",
        "Температура внутри помещения, °С",
        "Средняя температура наружного воздуха за отопительный период, °С"
    ]
    
    make_audit_table(
        doc,
        headers,
        rows,
        numeric_columns=[1, 2, 3, 4, 5, 6],
        col_widths=[Cm(4), Cm(3), Cm(4), Cm(4), Cm(3), Cm(3), Cm(4)]
    )
    doc.add_paragraph()
    
    portrait_section = doc.add_section(WD_SECTION.NEW_PAGE)
    portrait_section.orientation = WD_ORIENT.PORTRAIT
    p_width, p_height = portrait_section.page_height, portrait_section.page_width
    portrait_section.page_width = p_width
    portrait_section.page_height = p_height


def _section_per_building_transmission(doc, b: Dict[str, Any]) -> None:
    add_body_paragraph(
        doc,
        f"Ниже представлен расчёт трансмиссионных теплопотерь через наружные стены для здания {b['name']}.",
    )
    add_body_paragraph(doc, "Исходные данные:", first_line_indent=False)
    _bullet_paragraph(doc, f"Площадь наружных стен: S = {fmt_num(b['area_m2'], 2)} м²")
    _bullet_rich(doc, "Средняя наружная температура: ", ("t", "н"), f" = {fmt_num(b['t_outside_avg'], 1)} °С")
    _bullet_rich(doc, "Внутренняя температура: ", ("t", "в"), f" = {fmt_num(b['t_inside'], 1)} °С")
    _bullet_paragraph(doc, f"Продолжительность отопительного периода: z = {fmt_num(b['period_days'], 0)} сут.")
    _bullet_rich(doc, "Сопротивление теплопередаче до утепления: ", ("R", "до"), f" = {fmt_num(b['r_before'], 2)} м²·°С/Вт")
    _bullet_rich(doc, "Сопротивление теплопередаче после утепления: ", ("R", "после"), f" = {fmt_num(b['r_after'], 2)} м²·°С/Вт")
    
    add_body_paragraph(doc, "Теплопотери до мероприятия:")
    doc.add_paragraph()
    _add_equation(
        doc,
        eq.eq_q_kwh(
            subscript="1",
            r_subscript="до",
            area=fmt_num(b["area_m2"], 2),
            r_val=fmt_num(b["r_before"], 2),
            t_in=fmt_num(b["t_inside"], 1),
            t_out_avg=fmt_num(b["t_outside_avg"], 1),
            period_days=fmt_num(b["period_days"], 0),
            result_kwh=fmt_num(b["q1_kwh"], 1),
        ),
    )
    doc.add_paragraph()

    add_body_paragraph(doc, "Теплопотери после мероприятия:")
    doc.add_paragraph()
    _add_equation(
        doc,
        eq.eq_q_kwh(
            subscript="2",
            r_subscript="после",
            area=fmt_num(b["area_m2"], 2),
            r_val=fmt_num(b["r_after"], 2),
            t_in=fmt_num(b["t_inside"], 1),
            t_out_avg=fmt_num(b["t_outside_avg"], 1),
            period_days=fmt_num(b["period_days"], 0),
            result_kwh=fmt_num(b["q2_kwh"], 1),
        ),
    )
    doc.add_paragraph()

    add_body_paragraph(doc, "Экономия тепловой энергии, Гкал:")
    doc.add_paragraph()
    _add_equation(
        doc,
        eq.eq_delta_q(
            q1=fmt_num(b["q1_kwh"], 1),
            q2=fmt_num(b["q2_kwh"], 1),
            result_gcal=fmt_num(b["q_transmission_gcal"], 2),
        ),
    )
    doc.add_paragraph()


def _section_totals(doc, payload: Dict[str, Any], results: Dict[str, Any]) -> None:
    totals = results["totals"]
    shared = payload["shared"]
    add_body_paragraph(doc, f"ИТОГО: ΣQ = {fmt_num(totals['q_total_gcal'], 2)} Гкал", first_line_indent=False)

    fuel_type = shared.get("fuel_type", "gcal")
    meta = FUEL_META.get(fuel_type, FUEL_META["gcal"])

    if fuel_type != "gcal":
        add_body_paragraph(
            doc,
            f"Перевод потребленных Гкал в объем {meta['label']} ({meta['fuel_unit']}); удельная теплота сгорания принята за "
            f"{fmt_num(shared['fuel_calorific'], 4)} {meta['cv_unit']}:",
        )
        doc.add_paragraph()
        _add_equation(
            doc,
            eq.eq_fuel_savings(
                label=meta["label_short"],
                sum_q=fmt_num(totals["q_total_gcal"], 2),
                calorific=fmt_num(shared["fuel_calorific"], 4),
                result=fmt_num(totals["fuel_savings"], 3),
                unit=meta["fuel_unit"],
            ),
        )
        doc.add_paragraph()

    add_body_paragraph(doc, "Экономия тепловой энергии в денежном выражении, тг:")
    doc.add_paragraph()

    if fuel_type != "gcal":
        _add_equation(
            doc,
            eq.eq_money_savings(
                fuel_val=fmt_num(totals["fuel_savings"], 3),
                tariff=fmt_num(shared["fuel_tariff"], 2),
                multiplier=meta["multiplier"],
                result=fmt_money(totals["money_savings_tg"]),
            ),
        )
        doc.add_paragraph()
        add_body_paragraph(
            doc,
            f"где Cтэ = {fmt_num(shared['fuel_tariff'], 2)} {meta['tariff_unit']} — средняя цена {meta['label']}.",
        )
    else:
        _add_equation(
            doc,
            eq.eq_money_savings(
                fuel_val=fmt_num(totals["q_total_gcal"], 2),
                tariff=fmt_num(shared["fuel_tariff"], 2),
                multiplier="1",
                result=fmt_money(totals["money_savings_tg"]),
            ),
        )
        doc.add_paragraph()
        add_body_paragraph(
            doc,
            f"где Cтэ = {fmt_num(shared['fuel_tariff'], 2)} тг/Гкал — средняя цена тепловой энергии.",
        )

    invest = results["investment"]
    add_body_paragraph(doc, "Стоимость затрат на внедрение мероприятия, тг:")
    add_body_paragraph(
        doc,
        f"I = {fmt_money(invest['total_tg'])} тг (см. таблицу 3.3.7.2).",
        first_line_indent=False,
    )
    add_body_paragraph(doc, "Простой срок окупаемости мероприятия, год:")
    doc.add_paragraph()
    _add_equation(
        doc,
        eq.eq_pbp(
            investment=fmt_money(invest["total_tg"]),
            money=fmt_money(totals["money_savings_tg"]),
            result=fmt_years(results["finance"]["pbp_years"]),
        ),
    )
    doc.add_page_break()


def _section_table_3(doc, results: Dict[str, Any]) -> None:
    doc.add_paragraph()
    p_cap = doc.add_paragraph()
    p_cap.paragraph_format.first_line_indent = None
    p_cap.paragraph_format.space_before = Pt(6)
    p_cap.paragraph_format.space_after = Pt(0)
    p_cap.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    
    add_runs_with_highlight(p_cap, "Таблица 3.3.7.2 - Смета на утепление наружных стен", italic=True)

    rows = []
    for item in results["investment"]["items"]:
        rows.append(
            (
                item["name"],
                fmt_num(item["quantity"], 2),
                fmt_num(item["price_per_unit"], 2),
                fmt_money(item["cost_tg"]),
            )
        )
    rows.append(("Итого", "—", "—", fmt_money(results['investment']['total_tg'])))
    make_audit_table(
        doc,
        ["Наименование материалов/работ", "Количество, м²", "Цена за ед., тг", "Стоимость, тг"],
        rows,
        numeric_columns=[1, 2, 3], col_widths=[Cm(6), Cm(3), Cm(4), Cm(4)]
    )
    doc.add_paragraph()


def _section_table_4(doc, results: Dict[str, Any]) -> None:
    p_cap = doc.add_paragraph()
    p_cap.paragraph_format.first_line_indent = None
    p_cap.paragraph_format.space_before = Pt(6)
    p_cap.paragraph_format.space_after = Pt(0)
    p_cap.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    
    add_runs_with_highlight(p_cap, "Таблица 3.3.7.3 – Сводная экономия по зданиям", italic=True)

    rows = []
    for b in results["buildings"]:
        rows.append(
            (
                b["name"],
                fmt_num(b["q_total_gcal"], 2),
                fmt_money(b["money_savings_tg"]),
            )
        )
    rows.append(("ИТОГО", fmt_num(results["totals"]["q_total_gcal"], 2), fmt_money(results["totals"]["money_savings_tg"])))
    make_audit_table(
        doc,
        ["Наименование объекта", "Суммарная экономия, Гкал", "Экономия, тг/год"],
        rows,
        numeric_columns=[1, 2], col_widths=[Cm(9), Cm(4), Cm(4)]
    )
    doc.add_paragraph()


def _section_table_finance(doc, results: Dict[str, Any]) -> None:
    p_cap = doc.add_paragraph()
    p_cap.paragraph_format.first_line_indent = None
    p_cap.paragraph_format.space_before = Pt(6)
    p_cap.paragraph_format.space_after = Pt(0)
    p_cap.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    
    add_runs_with_highlight(p_cap, "Таблица 3.3.7.4 - Показатели эффективности проекта", italic=True)

    f = results["finance"]
    rows = [
        ("Ставка дисконтирования", "(%)", fmt_num(f["discount_rate"] * 100.0, 2)),
        ("Первоначальные инвестиции", "тг", fmt_money(results["investment"]["total_tg"])),
        ("Чистая приведённая стоимость (NPV)", "тг", fmt_money(f["npv_tg"])),
        ("Дисконтированный индекс доходности (DPI)", "о.е.", fmt_num(f["dpi"], 2)),
        (
            "Внутренняя норма доходности (IRR)",
            "(%)",
            fmt_percent(f["irr"]) if f["irr"] is not None else "—",
        ),
        ("Простой срок окупаемости (PBP)", "лет", fmt_years(f["pbp_years"])),
        ("Дисконтированный срок окупаемости (DPBP)", "год", fmt_years(f["dpbp_years"])),
    ]
    make_audit_table(doc, ["Наименование", "Ед. изм-я", "Значение"], rows, 
    numeric_columns=[2], col_widths=[Cm(11), Cm(3), Cm(3)])


def _section_chart(doc, results: Dict[str, Any]) -> None:
    new_section = doc.add_section(WD_SECTION.NEW_PAGE)
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_width, new_height = new_section.page_height, new_section.page_width
    new_section.page_width = new_width
    new_section.page_height = new_height

    png = render_npv_chart(results["finance"]["npv_by_year"])
    
    # Вставляем рисунок
    p_img = doc.add_paragraph()
    p_img.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p_img.paragraph_format.space_before = Pt(12)
    
    img_stream = io.BytesIO(png)
    p_img.add_run().add_picture(img_stream, width=Inches(9.0))
    
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p_cap.paragraph_format.space_before = Pt(6)
    p_cap.paragraph_format.space_after = Pt(12)
    
    add_runs_with_highlight(p_cap, "Рисунок 3.3.7.1 - Динамика изменения чистого дисконтированного дохода (NPV)", italic=True)
