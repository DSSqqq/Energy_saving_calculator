"""
Сборка раздела отчёта по модернизации АТП в формате .docx.
"""
from __future__ import annotations

import io
from pathlib import Path
from typing import Any, Dict

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm, Inches, Pt

from ..docx_utils.builder import (
    add_body_paragraph,
    add_caption,
    add_heading_para,
    add_runs_with_highlight,
    make_audit_table,
)
from ..docx_utils.chart import render_npv_chart
from ..docx_utils.number_fmt import (
    fmt_money,
    fmt_num,
    fmt_percent,
    fmt_years,
)
from . import equations as eq

# Переиспользуем base.docx из модуля окон
BASE_DOCX_PATH = Path(__file__).parent.parent / "windows" / "template" / "base.docx"


def build_docx(payload: Dict[str, Any], results: Dict[str, Any]) -> bytes:
    """Сборка документа для АТП."""
    if BASE_DOCX_PATH.exists():
        doc = Document(str(BASE_DOCX_PATH))
    else:
        doc = Document()
        _apply_fallback_styles(doc)

    add_heading_para(doc, "3.2.6 Организация (модернизация) автоматизированного теплового пункта")

    add_body_paragraph(
        doc,
        "Обследуемые объекты имеют центральное теплоснабжение, без возможности "
        "регулирования подачи тепловой энергии. Предлагается организовать современный "
        "автоматизированный тепловой пункт с функциями авторегулирования подачи тепловой "
        "энергии в зависимости от погодных условий и с учетом нерабочего периода, "
        "т.е. праздничные, выходные дни. Данное мероприятие позволит снизить потребление тепловой энергии."
    )

    add_body_paragraph(
        doc,
        "Для поддержания требуемого температурного графика в системе отопления планируется "
        "настроить регуляторы на отопление с датчиками наружного и внутреннего воздуха. "
        "По соответствующей программе регулятор, например, может осуществлять понижение "
        "температуры теплоносителя в дневное солнечное время. Автоматизированное управление "
        "отопительной нагрузкой позволяет получить экономию в осенне-весенний период, "
        "когда распространенной проблемой является наличие перетопов, связанное с особенностями "
        "центрального качественного регулирования тепловой нагрузки на источниках теплоснабжения. "
        "Индивидуальный учет тепловой энергии эффективен тогда, когда потребитель имеет "
        "возможность регулировать расход тепла в зависимости от своих собственных потребностей."
    )

    add_body_paragraph(
        doc,
        "Дежурное отопление осуществляется путем частичного использования основной системы "
        "отопления; при наличии централизованного теплоснабжения дежурное отопление может "
        "осуществляться путем использования системы централизованной регулировки параметров "
        "теплоносителя. Тепловая мощность дежурного отопления определяется в соответствии "
        "с теплопотерями помещений при пониженной температуре воздуха в них."
    )

    add_body_paragraph(
        doc,
        "В соответствии со строительными нормами и правилами расчетная температура воздуха "
        "в общественных и жилых помещениях в холодный период года должна соответствовать +21 °С. "
        "При внедрении дежурного отопления планируется снижение температуры воздуха внутри "
        "помещений в нерабочее время до +18 °С (или другого заданного дежурного значения)."
    )

    # 1. Таблица 3.2.6.1 - Исходные данные
    _section_table_inputs(doc, results)

    # 2. Математические формулы для первого здания
    if results["buildings"]:
        first_b = results["buildings"][0]
        _section_building_math(doc, first_b)

        if len(results["buildings"]) > 1:
            add_body_paragraph(
                doc,
                "Аналогичным методом выполнен расчет для остальных зданий, где это целесообразно. "
                "Результаты вычислений приведены ниже."
            )
            doc.add_paragraph()

    # 3. Таблица результатов
    _section_table_results(doc, results)

    # 4. Таблица финансовой эффективности
    _section_table_finance(doc, results)

    # 5. График NPV
    _section_chart(doc, results)

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


# --- helpers ------------------------------------------------------------------

def _apply_fallback_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    pf = normal.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.0


def _add_equation(doc, omml_element) -> None:
    """Вставляет OMML формулу в документ."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.left_indent = Pt(0)
    p.paragraph_format.right_indent = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(2)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    if omml_element.tag.endswith("oMathPara"):
        m_pr = omml_element.find(qn("m:oMathParaPr"))
        if m_pr is None:
            m_pr = OxmlElement("m:oMathParaPr")
            omml_element.insert(0, m_pr)
        m_jc = m_pr.find(qn("m:jc"))
        if m_jc is None:
            m_jc = OxmlElement("m:jc")
            m_pr.append(m_jc)
        m_jc.set(qn("m:val"), "centerGroup")

    p._p.append(omml_element)


def _bullet_rich(doc, *segments) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn as _qn

    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(1.89)
    p.paragraph_format.first_line_indent = Cm(-0.63)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    def _add_run(text, subscript=False):
        run = p.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        if subscript:
            rpr = run._element.get_or_add_rPr()
            va = OxmlElement("w:vertAlign")
            va.set(_qn("w:val"), "subscript")
            rpr.append(va)

    for seg in segments:
        if isinstance(seg, str):
            _add_run(seg)
        elif isinstance(seg, tuple) and len(seg) == 2:
            base, sub = seg
            _add_run(base)
            _add_run(sub, subscript=True)
        else:
            _add_run(str(seg))


# --- Секции отчета -----------------------------------------------------------

def _section_table_inputs(doc: Document, results: Dict[str, Any]) -> None:
    # Исходные данные по всем объектам в виде красивой таблицы (ландшафт)
    landscape_section = doc.add_section(WD_SECTION.NEW_PAGE)
    landscape_section.orientation = WD_ORIENT.LANDSCAPE
    w, h = landscape_section.page_height, landscape_section.page_width
    landscape_section.page_width = w
    landscape_section.page_height = h

    # Подпись
    p_cap = doc.add_paragraph()
    p_cap.paragraph_format.first_line_indent = None
    p_cap.paragraph_format.space_before = Pt(6)
    p_cap.paragraph_format.space_after = Pt(0)
    p_cap.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    add_runs_with_highlight(p_cap, "Таблица 3.2.6.1 - Исходные данные для модернизации АТП", italic=True)

    headers = [
        "Объект",
        "Q баз, Гкал/год",
        "t вн, °С",
        "t д, °С",
        "t ср.нар, °С",
        "t нар, °С",
        "N от, сут.",
        "N в.от, сут.",
        "День, ч",
        "Тариф, тг/Гкал",
    ]

    rows = []
    for b in results["buildings"]:
        rows.append(
            (
                b["name"],
                fmt_num(b["q_annual"], 2),
                fmt_num(b["t_inside"], 1),
                fmt_num(b["t_standby"], 1),
                fmt_num(b["t_outside_avg"], 1),
                fmt_num(b["t_outside_design"], 1),
                fmt_num(b["period_days"], 0),
                fmt_num(b["weekend_days"], 0),
                fmt_num(b["day_hours"], 0),
                fmt_money(b["fuel_tariff"]),
            )
        )

    make_audit_table(
        doc,
        headers,
        rows,
        numeric_columns=[1, 2, 3, 4, 5, 6, 7, 8, 9],
        col_widths=[Cm(4), Cm(2.5), Cm(2), Cm(2), Cm(2), Cm(2), Cm(2), Cm(2), Cm(2), Cm(3.5)]
    )

    doc.add_paragraph()

    # Возврат в книжную
    portrait_section = doc.add_section(WD_SECTION.NEW_PAGE)
    portrait_section.orientation = WD_ORIENT.PORTRAIT
    w_p, h_p = portrait_section.page_height, portrait_section.page_width
    portrait_section.page_width = w_p
    portrait_section.page_height = h_p


def _section_building_math(doc: Document, b: Dict[str, Any]) -> None:
    add_body_paragraph(
        doc,
        f"Выполним расчет экономии тепловой энергии на примере здания: {b['name']}.",
    )
    add_body_paragraph(doc, "Параметры объекта:", first_line_indent=False)
    _bullet_rich(doc, "Годовое базовое теплопотребление: ", ("Q", "баз"), f" = {fmt_num(b['q_annual'], 3)} Гкал")
    _bullet_rich(doc, "Температура внутри здания в рабочее время: ", ("t", "вн"), f" = {fmt_num(b['t_inside'], 1)} °С")
    _bullet_rich(doc, "Температура дежурного режима отопления: ", ("t", "д"), f" = {fmt_num(b['t_standby'], 1)} °С")
    _bullet_rich(doc, "Средняя наружная температура за отопительный период: ", ("t", "ср.нар"), f" = {fmt_num(b['t_outside_avg'], 1)} °С")
    _bullet_rich(doc, "Температура наиболее холодной пятидневки: ", ("t", "нар"), f" = {fmt_num(b['t_outside_design'], 1)} °С")
    _bullet_rich(doc, "Период отопления: ", ("N", "от"), f" = {fmt_num(b['period_days'], 0)} сут.")
    _bullet_rich(doc, "Количество рабочих дней: ", ("N", "р.от"), f" = {fmt_num(b['work_days'], 0)} сут.")
    _bullet_rich(doc, "Количество выходных и праздников: ", ("N", "в.от"), f" = {fmt_num(b['weekend_days'], 0)} сут.")
    _bullet_rich(doc, "Время поддержания нормированной температуры: ", ("τ", "день"), f" = {fmt_num(b['day_hours'], 0)} ч.")
    _bullet_rich(doc, "Время поддержания дежурной температуры в рабочие дни: ", ("τ", "ночь"), f" = {fmt_num(b['night_hours'], 0)} ч.")

    # 1. q0_max
    add_body_paragraph(doc, "Фактическая часовая тепловая нагрузка на отопление здания составляет, Гкал/ч:")
    doc.add_paragraph()
    _add_equation(
        doc,
        eq.eq_q0_max(
            q_base=fmt_num(b["q_annual"], 3),
            period_days=fmt_num(b["period_days"], 0),
            t_inside=fmt_num(b["t_inside"], 1),
            t_outside_avg=fmt_num(b["t_outside_avg"], 1),
            t_outside_design=fmt_num(b["t_outside_design"], 1),
            result_val=fmt_num(b["q0_max"], 6),
        ),
    )
    doc.add_paragraph()

    # 2. Q_p
    add_body_paragraph(
        doc,
        "Расход тепловой энергии в рабочее (дневное) время отопительного периода, Гкал:"
    )
    doc.add_paragraph()
    _add_equation(
        doc,
        eq.eq_q_p(
            work_days=fmt_num(b["work_days"], 0),
            day_hours=fmt_num(b["day_hours"], 0),
            t_inside=fmt_num(b["t_inside"], 1),
            t_outside_avg=fmt_num(b["t_outside_avg"], 1),
            t_outside_design=fmt_num(b["t_outside_design"], 1),
            q0_max=fmt_num(b["q0_max"], 6),
            result_val=fmt_num(b["q_p"], 3),
        ),
    )
    doc.add_paragraph()

    # 3. Q_v
    add_body_paragraph(
        doc,
        "Расход тепловой энергии в дежурное (ночное и выходное) время отопительного периода, Гкал:"
    )
    doc.add_paragraph()
    _add_equation(
        doc,
        eq.eq_q_v(
            work_days=fmt_num(b["work_days"], 0),
            night_hours=fmt_num(b["night_hours"], 0),
            weekend_days=fmt_num(b["weekend_days"], 0),
            t_standby=fmt_num(b["t_standby"], 1),
            t_outside_avg=fmt_num(b["t_outside_avg"], 1),
            t_outside_design=fmt_num(b["t_outside_design"], 1),
            q0_max=fmt_num(b["q0_max"], 6),
            result_val=fmt_num(b["q_v"], 3),
        ),
    )
    doc.add_paragraph()

    # 4. Q_ob
    add_body_paragraph(
        doc,
        "Общее потребление тепловой энергии от внедрения дежурного отопления за отопительный период, Гкал:"
    )
    doc.add_paragraph()
    _add_equation(
        doc,
        eq.eq_q_ob(
            q_p=fmt_num(b["q_p"], 3),
            q_v=fmt_num(b["q_v"], 3),
            result_val=fmt_num(b["q_ob"], 3),
        ),
    )
    doc.add_paragraph()

    # 5. Q_new & delta_Q
    add_body_paragraph(
        doc,
        "Общая экономия тепловой энергии за счет организации автоматизированного теплового пункта, "
        "при учете дополнительного снижения теплопотребления на 9 % за счет устранения перетопов "
        "в осенне-весенний период, Гкал/год:"
    )
    doc.add_paragraph()
    _add_equation(
        doc,
        eq.eq_delta_q(
            q_base=fmt_num(b["q_annual"], 3),
            q_ob=fmt_num(b["q_ob"], 3),
            result_new=fmt_num(b["q_new"], 3),
            result_delta=fmt_num(b["delta_q"], 3),
        ),
    )
    doc.add_paragraph()

    # 6. Money Savings
    add_body_paragraph(doc, "Годовая экономия в денежном выражении составит, тг:")
    doc.add_paragraph()
    _add_equation(
        doc,
        eq.eq_money_savings(
            delta_q=fmt_num(b["delta_q"], 3),
            tariff=fmt_num(b["fuel_tariff"], 2),
            result_val=fmt_money(b["money_savings_tg"]),
        ),
    )
    doc.add_paragraph()

    # 7. Payback
    add_body_paragraph(doc, "Простой срок окупаемости установки автоматики для здания составит, лет:")
    doc.add_paragraph()
    _add_equation(
        doc,
        eq.eq_pbp_atp(
            investment=fmt_money(b["investment"]),
            savings=fmt_money(b["money_savings_tg"]),
            result_val=fmt_years(b["pbp_years"]),
        ),
    )
    doc.add_page_break()


def _section_table_results(doc: Document, results: Dict[str, Any]) -> None:
    p_cap = doc.add_paragraph()
    p_cap.paragraph_format.first_line_indent = None
    p_cap.paragraph_format.space_before = Pt(6)
    p_cap.paragraph_format.space_after = Pt(0)
    p_cap.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    add_runs_with_highlight(
        p_cap,
        "Таблица 3.2.6.2 - Показатели энергоэффективности по модернизации АТП",
        italic=True,
    )

    headers = [
        "№",
        "Объект",
        "Энергопотребление баз., Гкал",
        "Экономия в натур. выр., Гкал",
        "Экономия в денежном выр., тг",
        "Затраты, тг",
        "Срок окупаемости, лет",
    ]

    rows = []
    for idx, b in enumerate(results["buildings"], 1):
        rows.append(
            (
                str(idx),
                b["name"],
                fmt_num(b["q_annual"], 2),
                fmt_num(b["delta_q"], 3),
                fmt_money(b["money_savings_tg"]),
                fmt_money(b["investment"]),
                fmt_years(b["pbp_years"]),
            )
        )

    rows.append(
        (
            "Итого",
            "Итого",
            fmt_num(results["totals"]["q_base_gcal"], 2),
            fmt_num(results["totals"]["q_total_gcal"], 3),
            fmt_money(results["totals"]["money_savings_tg"]),
            fmt_money(results["totals"]["investment_total_tg"]),
            fmt_years(results["finance"]["pbp_years"]),
        )
    )

    make_audit_table(
        doc,
        headers,
        rows,
        numeric_columns=[2, 3, 4, 5, 6],
        col_widths=[Cm(1), Cm(4), Cm(3), Cm(3), Cm(3), Cm(3), Cm(2.5)]
    )
    doc.add_paragraph()


def _section_table_finance(doc: Document, results: Dict[str, Any]) -> None:
    p_cap = doc.add_paragraph()
    p_cap.paragraph_format.first_line_indent = None
    p_cap.paragraph_format.space_before = Pt(6)
    p_cap.paragraph_format.space_after = Pt(0)
    p_cap.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    add_runs_with_highlight(p_cap, "Таблица 3.2.6.3 - Показатели эффективности проекта", italic=True)

    f = results["finance"]
    rows = [
        ("Ставка дисконтирования", "(%)", fmt_num(f["discount_rate"] * 100.0, 2)),
        ("Первоначальные инвестиции", "тг", fmt_money(results["investment"]["total_tg"])),
        ("Чистая приведённая стоимость (NPV)", "тг", fmt_money(f["npv_tg"])),
        ("Дисконтированный индекс доходности (DPI)", "о.е.", fmt_num(f["dpi"], 2)),
        (
            "Внутренняя норма доходности (IRR)",
            "(%)",
            fmt_percent(f["irr"]) if f["irr"] is not None else "—",
        ),
        ("Простой срок окупаемости (PBP)", "лет", fmt_years(f["pbp_years"])),
        ("Дисконтированный срок окупаемости (DPBP)", "год", fmt_years(f["dpbp_years"])),
    ]

    make_audit_table(
        doc,
        ["Наименование", "Ед. изм.", "Значение"],
        rows,
        numeric_columns=[2],
        col_widths=[Cm(11), Cm(3), Cm(3)]
    )
    doc.add_paragraph()


def _section_chart(doc: Document, results: Dict[str, Any]) -> None:
    # Добавляем раздел альбомного формата для графика
    new_section = doc.add_section(WD_SECTION.NEW_PAGE)
    new_section.orientation = WD_ORIENT.LANDSCAPE
    w, h = new_section.page_height, new_section.page_width
    new_section.page_width = w
    new_section.page_height = h

    png = render_npv_chart(results["finance"]["npv_by_year"])

    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    run.add_picture(io.BytesIO(png), width=Inches(9.0))

    p_cap = doc.add_paragraph()
    p_cap.paragraph_format.first_line_indent = None
    p_cap.paragraph_format.space_before = Pt(6)
    p_cap.paragraph_format.space_after = Pt(12)
    p_cap.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    add_runs_with_highlight(
        p_cap,
        "Диаграмма 3.2.6.1 – Чистая приведенная стоимость проекта по модернизации теплового пункта",
        italic=True,
    )
